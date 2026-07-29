from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree


SKILL_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = SKILL_ROOT / "assets" / "response_letter_template.docx"

FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
BLUE = RGBColor(68, 114, 196)
RED = RGBColor(192, 0, 0)
BOX_COLOR = "4472C4"
HEADER_FILL = "B4C7E7"
COMMENT_HEADER_FILL = "B4C6E7"
COMMENT_BODY_FILL = "ECF1F8"
COMMENT_TABLE_STYLE = "Grid Table 4 Accent 1"


def set_font(run, *, size=11, bold=False, italic=False, color=BLACK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), FONT)
    return run


def add_run(paragraph, text, **formatting):
    return set_font(paragraph.add_run(text), **formatting)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)


def set_table_borders(table, size="8"):
    borders = OxmlElement("w:tblBorders")
    table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), BOX_COLOR)
        borders.append(border)


def add_comment_table(document, label, comment_placeholder):
    table = document.add_table(rows=2, cols=1)
    table.style = COMMENT_TABLE_STYLE
    section = document.sections[-1]
    width = str(int((section.page_width - section.left_margin - section.right_margin) / 635))
    tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table._tbl.tblPr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), width)
    tbl_w.set(qn("w:type"), "dxa")
    for grid_col in table._tbl.tblGrid.findall(qn("w:gridCol")):
        grid_col.set(qn("w:w"), width)
    for row in table.rows:
        for cell in row.cells:
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().insert(0, tc_w)
            tc_w.set(qn("w:w"), width)
            tc_w.set(qn("w:type"), "dxa")

    header = table.cell(0, 0)
    shade_cell(header, COMMENT_HEADER_FILL)
    header.paragraphs[0].clear()
    add_run(header.paragraphs[0], label, bold=True)

    body = table.cell(1, 0)
    shade_cell(body, COMMENT_BODY_FILL)
    body.paragraphs[0].clear()
    add_run(body.paragraphs[0], comment_placeholder)
    return table


def add_divider(document):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "dashed")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BOX_COLOR)
    borders.append(bottom)
    p_pr.append(borders)


def add_page_number(document):
    paragraph = document.sections[0].footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.clear()
    run = add_run(paragraph, "Page ", size=9)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)


def configure_styles(document):
    for style in document.styles:
        if not hasattr(style, "font"):
            continue
        style.font.name = FONT
        r_pr = style._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{attribute}"), FONT)
    normal = document.styles["Normal"]
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)


def ensure_comment_table_style(document):
    try:
        document.styles[COMMENT_TABLE_STYLE]
        return
    except KeyError:
        pass
    xml = '''<w:style xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="table" w:styleId="GridTable4Accent1">
      <w:name w:val="Grid Table 4 Accent 1"/><w:basedOn w:val="TableNormal"/><w:uiPriority w:val="49"/>
      <w:tblPr><w:tblStyleRowBandSize w:val="1"/><w:tblStyleColBandSize w:val="1"/><w:tblBorders>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="8EAADB"/><w:left w:val="single" w:sz="4" w:space="0" w:color="8EAADB"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="8EAADB"/><w:right w:val="single" w:sz="4" w:space="0" w:color="8EAADB"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="8EAADB"/><w:insideV w:val="single" w:sz="4" w:space="0" w:color="8EAADB"/>
      </w:tblBorders></w:tblPr>
      <w:tblStylePr w:type="firstRow"><w:rPr><w:b/><w:bCs/></w:rPr><w:tcPr><w:tcBorders>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="4472C4"/><w:left w:val="single" w:sz="4" w:space="0" w:color="4472C4"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="4472C4"/><w:right w:val="single" w:sz="4" w:space="0" w:color="4472C4"/>
      </w:tcBorders></w:tcPr></w:tblStylePr>
      <w:tblStylePr w:type="band1Horz"><w:tcPr><w:shd w:val="clear" w:color="auto" w:fill="D9E2F3"/></w:tcPr></w:tblStylePr>
    </w:style>'''
    document.styles.element.append(etree.fromstring(xml.encode("utf-8")))


def build_template(output_path=DEFAULT_OUTPUT):
    document = Document()
    configure_styles(document)
    ensure_comment_table_style(document)
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "Response Letter", size=16, bold=True)

    paragraph = document.add_paragraph()
    add_run(paragraph, "Manuscript Number: ", bold=True)
    add_run(paragraph, "[Manuscript ID]")
    paragraph = document.add_paragraph()
    add_run(paragraph, "Title: ", bold=True)
    add_run(paragraph, "[Manuscript title]")

    add_divider(document)
    paragraph = document.add_paragraph()
    add_run(paragraph, "Dear Editor,", bold=True)
    paragraph = document.add_paragraph()
    add_run(paragraph, "[Insert the concise overall response after all point-by-point replies are finalized.]")
    add_comment_table(document, "Editor Comments", "[Paste the Editor's comments verbatim here.]")
    paragraph = document.add_paragraph()
    add_run(paragraph, "Response: ", bold=True)
    add_run(paragraph, "[Insert the direct response to the Editor here.]")

    add_divider(document)
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(heading, "Response to Reviewer 1", size=12, bold=True)
    add_comment_table(document, "Comment 1", "[Paste Reviewer 1's original comment verbatim here.]")

    paragraph = document.add_paragraph()
    add_run(paragraph, "Response:", bold=True)
    paragraph = document.add_paragraph()
    add_run(paragraph, "[Give the direct answer first, followed by the minimum evidence and manuscript action.]")

    paragraph = document.add_paragraph()
    add_run(paragraph, "Revised text in the manuscript:", bold=True, color=BLUE)
    paragraph = document.add_paragraph()
    add_run(paragraph, "[Paste the complete revised context in blue. ", color=BLUE)
    add_run(paragraph, "Bold only the words actually changed", bold=True, color=BLUE)
    add_run(paragraph, ".]", color=BLUE)
    paragraph = document.add_paragraph()
    add_run(paragraph, "(see Section X, Fig. X, Table X, or Lines X-X in the revised manuscript)", italic=True, color=RED)

    add_divider(document)
    paragraph = document.add_paragraph()
    add_run(paragraph, "Added reference:", italic=True, color=BLUE)
    paragraph = document.add_paragraph()
    add_run(paragraph, "[Insert the complete added reference here.]", color=BLUE)

    add_divider(document)
    paragraph = document.add_paragraph()
    add_run(paragraph, "Response-only figure or table:", italic=True)
    placeholder = document.add_table(rows=1, cols=1)
    set_table_borders(placeholder)
    cell = placeholder.cell(0, 0)
    set_cell_margins(cell, top=500, bottom=500)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cell.paragraphs[0], "[Insert response-only figure or table here]", italic=True)
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(caption, "Fig. R1. [Insert a self-contained caption.]", bold=True)
    paragraph = document.add_paragraph()
    add_run(paragraph, "(response-only evidence; state whether it was added to the manuscript)", italic=True, color=RED)

    add_page_number(document)
    document.core_properties.title = "Response Letter Template"
    document.core_properties.subject = "De-identified reusable response-letter template"
    document.core_properties.author = ""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


if __name__ == "__main__":
    print(build_template())
